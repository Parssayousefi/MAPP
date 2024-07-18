import os
import re
from docx import Document
import pandas as pd
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows

def sanitize_filename(filename):
    """
    Sanitize the block name to be used as a valid filename by removing or replacing invalid characters.
    """
    return re.sub(r'[\\/*?:"<>|]', "", filename)

def create_excel_file(data, output_path):
    """
    Create an Excel file with items, content, and questionnaire name.
    
    :param data: List of dictionaries containing 'item', 'content', and 'questionnaire'
    :param output_path: Path to save the Excel file
    """
    df = pd.DataFrame(data)
    
    # Create a new workbook and select the active sheet
    wb = Workbook()
    ws = wb.active
    ws.title = "Questionnaire Items"

    # Write the DataFrame to the worksheet
    for r in dataframe_to_rows(df, index=False, header=True):
        ws.append(r)

    # Save the workbook
    excel_path = os.path.join(output_path, "questionnaire_items.xlsx")
    wb.save(excel_path)
    print(f"Excel file created: {excel_path}")

def parse_and_create_docs(docx_path, output_path):
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Error opening the document: {e}")
        return

    current_block_name = ""
    block_started = False
    new_doc = None
    excel_data = []
    current_item = ""
    current_content = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Start of Block:"):
            if new_doc:  # Save and close the previous document before starting a new one
                filename = os.path.join(output_path, sanitize_filename(current_block_name) + ".docx")
                new_doc.save(filename)
                # Add the last item of the previous block to excel_data
                if current_item:
                    excel_data.append({
                        'item': current_item,
                        'content': current_content.strip(),
                        'questionnaire': current_block_name
                    })
            current_block_name = text[len("Start of Block:"):].strip()
            new_doc = Document()  # Start a new document
            block_started = True
            current_item = ""
            current_content = ""
        elif text.startswith("End of Block:"):
            if new_doc:  # End of block, save the document
                filename = os.path.join(output_path, sanitize_filename(current_block_name) + ".docx")
                new_doc.save(filename)
                new_doc = None
            block_started = False
            # Add the last item of the block to excel_data
            if current_item:
                excel_data.append({
                    'item': current_item,
                    'content': current_content.strip(),
                    'questionnaire': current_block_name
                })
        else:
            if block_started:  # If we are currently in a block, add the paragraph to the new doc
                new_doc.add_paragraph(text)
                if text.endswith('?') or text.endswith(':'):  # Assuming questions end with '?' or ':'
                    if current_item:  # Save the previous item if exists
                        excel_data.append({
                            'item': current_item,
                            'content': current_content.strip(),
                            'questionnaire': current_block_name
                        })
                    current_item = text
                    current_content = ""
                else:
                    current_content += f" {text}"

    # Handle tables separately
    if new_doc:  # If there's an open document, save it
        filename = os.path.join(output_path, sanitize_filename(current_block_name) + ".docx")
        new_doc.save(filename)

    # Now, process tables
    for table in doc.tables:
        if new_doc:  # Assuming tables should be added to the end of the current document
            new_table = new_doc.add_table(rows=0, cols=0)  # Create a new table to copy the content
            for row in table.rows:
                new_cells = new_table.add_row().cells
                for index, cell in enumerate(row.cells):
                    if index < len(new_cells):  # Prevent index error
                        new_cells[index].text = cell.text
            filename = os.path.join(output_path, sanitize_filename(current_block_name) + ".docx")
            new_doc.save(filename)

    # Create Excel file
    create_excel_file(excel_data, output_path)

# Example usage
docx_path = r"C:\Users\prsyu\OneDrive\Bidlung\University\M.S. Leiden University\M.S. Neuroscience (Research)\MAPP\Analysis\Github_Clone\MAPP\Questionairs\Scale_Survey.docx"
output_path = r"C:\Users\prsyu\OneDrive\Bidlung\University\M.S. Leiden University\M.S. Neuroscience (Research)\MAPP\Analysis\Github_Clone\MAPP\Questionairs\word_files"

# Ensure output directory exists
os.makedirs(output_path, exist_ok=True)

# Now call your function
parse_and_create_docs(docx_path, output_path)